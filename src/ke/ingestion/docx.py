"""DOCX file ingestor using python-docx."""

from __future__ import annotations

import time
from pathlib import Path

from ke.core.models import IngestionResult
from ke.ingestion.base import Ingestor, chunk_text, create_entries_from_chunks


class DocxIngestor(Ingestor):
    """Ingest Word DOCX files."""

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def ingest(self, path: Path) -> IngestionResult:
        start = time.time()
        result = IngestionResult(source_path=str(path))

        try:
            from docx import Document

            doc = Document(str(path))

            paragraphs = []
            headings = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading", "").strip() or "1"
                    headings.append(f"H{level}: {text}")
                    paragraphs.append(f"## {text}")
                else:
                    paragraphs.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            tags = []
            if headings:
                tags.extend([h.split(": ", 1)[1] for h in headings[:5]])

            combined_text = "\n\n".join(paragraphs)
            chunks = chunk_text(combined_text)
            entries = create_entries_from_chunks(
                chunks, path, "docx",
                metadata={"headings": headings, "paragraph_count": len(paragraphs)},
                tags=tags,
            )

            result.entries = entries
            result.chunks_created = len(chunks)

        except ImportError:
            result.errors.append("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            result.errors.append(f"Error reading DOCX: {e}")

        result.duration_ms = (time.time() - start) * 1000
        return result
